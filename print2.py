from docx import Document
import os
import sys
from docx2python import docx2python
import win32com.client as win32

def edit_word_doc(value_dict, value_dic):
    # Open the document template
    doc = Document('D:\Desktop Apps\medigine\out.docx')
    
    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in value_dict.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)
    
    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in value_dic.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
    
    # Save the edited document
    output_path = os.path.abspath('D:\Desktop Apps\medigine\out.docx')
    doc.save(output_path)
    
    # Print the edited document
    print_word_doc(output_path)

def print_word_doc(doc_path):
    # Create a Word application instance
    word_app = win32.Dispatch('Word.Application')
    
    # Open the document
    doc = word_app.Documents.Open(doc_path)
    
    # Print the document
    doc.PrintOut()
    
    # Close the document without saving
    doc.Close(False)
    
    # Quit the Word application
    word_app.Quit()

if __name__ == "__main__":
    # Create a dictionary from command line arguments
    value_dict = {
        'XYZ': sys.argv[1],
        '{MCHARGES}': sys.argv[2],
        '{SCHARGES}': sys.argv[3],
        '{TOTAL}': sys.argv[4],
        '{THANKS}': sys.argv[5],
     
    }
    value_dic = {
        '{PAYMENT}': sys.argv[6],
        '{AMTWORDS}': sys.argv[7]
    }
    
    # Edit and print the Word document
    edit_word_doc(value_dict, value_dic)



